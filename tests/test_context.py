"""Tests for load_personal_context and generate_personal_career_and_finance_plan.

Most docx fixtures are generated on the fly with python-docx so nothing binary
needs to be committed. tests/data/Career_Plan_Test.docx is the one exception: a
synthetic (fake numbers, fake content) but Word-shaped fixture, checked in under
a targeted carve-out in hooks/privacy-guard.sh - see that file for why it's safe
to commit where a real personal .docx never would be.
"""
import shutil
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from src.context import _heading_prefix, generate_personal_career_and_finance_plan, load_personal_context

FIXTURE_DOCX = Path(__file__).parent / "data" / "Career_Plan_Test.docx"


def make_docx(path: Path, paragraphs: list[str]) -> None:
    """Write a minimal real .docx file at `path` with one paragraph per string."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def make_docx_with_table(path: Path, table_rows: list[list[str]], heading: str | None = None) -> None:
    """Write a .docx with an optional Heading 1 paragraph followed by a table."""
    document = Document()
    if heading is not None:
        document.add_heading(heading, level=1)
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    for row_idx, row in enumerate(table_rows):
        for col_idx, cell_text in enumerate(row):
            table.cell(row_idx, col_idx).text = cell_text
    document.save(path)


class TestLoadPersonalContext:
    def test_missing_directory_returns_empty_string(self, tmp_path):
        missing = tmp_path / "does-not-exist"

        assert load_personal_context(missing) == ""

    def test_empty_directory_returns_empty_string(self, tmp_path):
        assert load_personal_context(tmp_path) == ""

    def test_reads_single_docx_file(self, tmp_path):
        make_docx(tmp_path / "career_plan.docx", ["Target: staff engineer by 2030.", "Risk appetite: moderate."])

        result = load_personal_context(tmp_path)

        assert "## career_plan" in result
        assert "Target: staff engineer by 2030." in result
        assert "Risk appetite: moderate." in result

    def test_concatenates_multiple_docx_files_in_sorted_order(self, tmp_path):
        make_docx(tmp_path / "b_career_plan.docx", ["career content"])
        make_docx(tmp_path / "a_cv.docx", ["cv content"])

        result = load_personal_context(tmp_path)

        assert result.index("a_cv") < result.index("b_career_plan")

    def test_ignores_non_docx_files(self, tmp_path):
        make_docx(tmp_path / "career_plan.docx", ["career content"])
        (tmp_path / "notes.txt").write_text("should be ignored")

        result = load_personal_context(tmp_path)

        assert "should be ignored" not in result

    def test_skips_blank_paragraphs(self, tmp_path):
        make_docx(tmp_path / "career_plan.docx", ["Real content.", "   ", ""])

        result = load_personal_context(tmp_path)

        assert result == "## career_plan\nReal content."

    def test_unreadable_docx_is_skipped_not_raised(self, tmp_path, capsys):
        corrupt = tmp_path / "corrupt.docx"
        corrupt.write_bytes(b"not a real docx file")
        make_docx(tmp_path / "career_plan.docx", ["career content"])

        result = load_personal_context(tmp_path)

        assert "career content" in result
        assert "could not read corrupt.docx" in capsys.readouterr().out


class TestLoadPersonalContextTables:
    """Regression coverage for the table-drop bug: document.paragraphs silently
    omits table content, so a career plan's salary tables never reached the
    system prompt even though extraction raised no error."""

    def test_table_is_rendered_as_markdown_pipe_table(self, tmp_path):
        make_docx_with_table(tmp_path / "career_plan.docx", [["Plan", "Salary"], ["A", "CHF 300k"]])

        result = load_personal_context(tmp_path)

        assert "| Plan | Salary |" in result
        assert "| --- | --- |" in result
        assert "| A | CHF 300k |" in result

    def test_table_stays_after_its_heading_in_document_order(self, tmp_path):
        make_docx_with_table(tmp_path / "career_plan.docx", [["Plan", "Salary"]], heading="Salary bands")

        result = load_personal_context(tmp_path)

        assert result.index("Salary bands") < result.index("| Plan | Salary |")

    def test_heading_style_maps_to_markdown_heading(self, tmp_path):
        document = Document()
        document.add_heading("Career plan", level=1)
        document.add_paragraph("Body text.")
        document.save(tmp_path / "career_plan.docx")

        result = load_personal_context(tmp_path)

        assert "### Career plan" in result
        assert "Body text." in result
        assert "### Body text." not in result


class TestHeadingPrefix:
    def test_missing_style_is_treated_as_body_text(self):
        """Some real-world paragraphs resolve `.style` to None (e.g. a style id
        the document's style part doesn't define) - must not raise AttributeError."""
        paragraph = SimpleNamespace(style=None)

        assert _heading_prefix(paragraph) is None

    def test_non_heading_style_is_treated_as_body_text(self):
        paragraph = SimpleNamespace(style=SimpleNamespace(name="Normal"))

        assert _heading_prefix(paragraph) is None

    def test_heading_1_maps_to_three_hashes(self):
        paragraph = SimpleNamespace(style=SimpleNamespace(name="Heading 1"))

        assert _heading_prefix(paragraph) == "###"


class TestLoadPersonalContextAgainstFixtureFile:
    """End-to-end completeness check against a Word-shaped fixture, not the
    minimal ones built paragraph-by-paragraph above. Reads every paragraph and
    table cell straight from the source .docx and asserts each one survives
    extraction verbatim - this is what would have caught the table-drop bug
    immediately, since the fixture carries a real salary table."""

    def test_every_paragraph_and_table_cell_survives_extraction(self, tmp_path):
        shutil.copy(FIXTURE_DOCX, tmp_path / FIXTURE_DOCX.name)
        source = Document(FIXTURE_DOCX)

        expected_snippets = [p.text.strip() for p in source.paragraphs if p.text.strip()]
        expected_snippets += [
            cell.text.strip()
            for table in source.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        assert expected_snippets, "fixture has no extractable content - test would pass vacuously"

        result = load_personal_context(tmp_path)

        missing = [snippet for snippet in expected_snippets if snippet not in result]
        assert missing == [], f"lost during extraction: {missing}"


class TestGeneratePersonalCareerAndFinancePlan:
    """Regression coverage for the missing `return` bug: the function built the
    prompt string but fell off the end without returning it, so callers silently
    got None and the agent ran with no system prompt at all."""

    def test_returns_the_prompt_as_a_string(self, monkeypatch):
        monkeypatch.setattr("src.context.load_personal_context", lambda: "")

        result = generate_personal_career_and_finance_plan()

        assert isinstance(result, str)
        assert "personal advisor for career and financial planning" in result

    def test_appends_personal_context_when_present(self, monkeypatch):
        monkeypatch.setattr("src.context.load_personal_context", lambda: "## cv\nSome career history.")

        result = generate_personal_career_and_finance_plan()

        assert "# Personal context" in result
        assert "Some career history." in result

    def test_omits_personal_context_section_when_absent(self, monkeypatch):
        monkeypatch.setattr("src.context.load_personal_context", lambda: "")

        result = generate_personal_career_and_finance_plan()

        assert "# Personal context" not in result
