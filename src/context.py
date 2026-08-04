"""Loads personal context (CV, career plan) from local .docx files for the system prompt."""
import re
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

PERSONAL_DATA_DIR = Path(__file__).resolve().parent.parent / "data" / "personal"

_HEADING_STYLE_RE = re.compile(r"^Heading (\d)$")


def _heading_prefix(paragraph: Paragraph) -> str | None:
    """
    Markdown heading prefix for a paragraph in "Heading N" style, nested one level
    below the file's own "## <filename>" section heading; None for body text.
    """
    style = paragraph.style
    match = _HEADING_STYLE_RE.match(style.name) if style is not None else None
    return "#" * (int(match.group(1)) + 2) if match else None


def _iter_block_items(document: DocxDocument):
    """
    Yield paragraphs and tables in document order. python-docx exposes them as
    separate `document.paragraphs` / `document.tables` collections with no
    positional relationship, so reading those in isolation loses which heading a
    table belongs under and drops tables to the end of the extracted text.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _render_table(table: Table) -> str:
    """Render a docx table as a markdown pipe table; models parse those reliably."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(row)]
    if not rows:
        return ""

    header, *body = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    lines += ["| " + " | ".join(row) + " |" for row in body]
    return "\n".join(lines)


def load_personal_context(directory: Path = PERSONAL_DATA_DIR) -> str:
    """
    Read every .docx file in `directory` and concatenate their content - paragraphs
    and tables interleaved in document order, headings mapped to markdown - so
    personal context can be refreshed by editing the files without touching any code.
    :param directory: folder containing personal .docx files (gitignored).
    :return: concatenated markdown from all readable .docx files found, one section
             per file; "" if the directory is missing or contains none.
    """
    if not directory.is_dir():
        return ""

    sections = []
    for docx_path in sorted(directory.glob("*.docx")):
        try:
            document = Document(docx_path)
        except (PackageNotFoundError, zipfile.BadZipFile) as e:
            print(f"Warning: could not read {docx_path.name}: {e}")
            continue

        blocks = []
        for item in _iter_block_items(document):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if not text:
                    continue
                prefix = _heading_prefix(item)
                blocks.append(f"{prefix} {text}" if prefix else text)
            else:
                table_text = _render_table(item)
                if table_text:
                    blocks.append(table_text)

        if blocks:
            sections.append(f"## {docx_path.stem}\n" + "\n\n".join(blocks))

    return "\n\n".join(sections)


def generate_personal_career_and_finance_plan() -> str:
    """Generates the system prompt for personal career and finance plan."""
    system_prompt = (
        f"Today's date is {date.today().isoformat()}. Use it for any date-relative reasoning "
        "(current age, years remaining to a milestone, time elapsed) - you have no other way to "
        "know the current date, and the career plan below states milestones as target years "
        "(\"age 49 by 2036\"), not a current age, so getting this wrong compounds into every "
        "date-relative answer.\n\n"
        "You are a personal advisor for career and financial planning, treated as one problem: "
        "the user's earnings trajectory determines what their portfolio can become, and their "
        "portfolio target determines what their career must deliver and by when. Always reason "
        "about the two together, not as separate topics.\n\n"
        "However, the user might want to deep dive into the financial planning alone. In this case "
        "you must be able to deep dive in the topic, understand the financial needs of the user, "
        "check if the financial plan makes sense and make suggestion."
        "When advising on the career plan, take the user's personal background and professional "
        "experience into account rather than giving generic career advice — proposals must be "
        "realistic given their actual trajectory, not a hypothetical blank-slate candidate.\n\n"
        "Be direct: state disagreement plainly, do not hedge, and do not restate the user's own "
        "numbers back to them. Ground every answer in their actual plan and figures rather than "
        "generic advice."
    )

    personal_context = load_personal_context()
    if personal_context:
        system_prompt += "\n\n# Personal context\n" + personal_context

    return system_prompt
